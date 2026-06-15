from dataclasses import dataclass, field
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from pypdf import PdfReader

from app.ingestion.cleaning import clean_page_texts, normalize_document_text
from app.ingestion.table_extractors import docx_table_to_markdown, extract_html_tables, extract_pdf_tables


SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    source: str
    metadata: dict


@dataclass(frozen=True)
class LoadResult:
    documents: list[LoadedDocument]
    skipped: list[str] = field(default_factory=list)
    errors: dict[str, str] = field(default_factory=dict)


def load_document(path: Path) -> list[LoadedDocument]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported document type: {suffix}")

    if suffix in {".txt", ".md"}:
        text = normalize_document_text(path.read_text(encoding="utf-8"))
        return [
            LoadedDocument(
                text=text,
                source=str(path),
                metadata={"file_type": suffix},
            )
        ] if text else []

    if suffix == ".docx":
        return _load_docx(path)

    if suffix in {".html", ".htm"}:
        return _load_html(path)

    return _load_pdf(path)


def load_documents_from_dir(directory: Path) -> LoadResult:
    documents: list[LoadedDocument] = []
    skipped: list[str] = []
    errors: dict[str, str] = {}

    if not directory.exists():
        return LoadResult(documents=[], skipped=[], errors={str(directory): "Directory does not exist"})

    for path in sorted(p for p in directory.rglob("*") if p.is_file()):
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            skipped.append(str(path))
            continue

        try:
            documents.extend(load_document(path))
        except Exception as exc:
            errors[str(path)] = str(exc)

    return LoadResult(documents=documents, skipped=skipped, errors=errors)


def _load_pdf(path: Path) -> list[LoadedDocument]:
    reader = PdfReader(str(path))
    documents: list[LoadedDocument] = []
    page_records: list[tuple[int, str]] = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_records.append((page_index, text))

    cleaned_pages = clean_page_texts([text for _, text in page_records])

    for (page_index, _), text in zip(page_records, cleaned_pages, strict=False):
        if not text:
            continue
        documents.append(
            LoadedDocument(
                text=text,
                source=str(path),
                metadata={"file_type": ".pdf", "page": page_index, "content_type": "text"},
            )
        )

    for table in extract_pdf_tables(path):
        metadata = {"file_type": ".pdf", **table.metadata}
        documents.append(LoadedDocument(text=table.markdown, source=str(path), metadata=metadata))

    return documents


def _load_docx(path: Path) -> list[LoadedDocument]:
    doc = DocxDocument(str(path))
    paragraphs: list[str] = []
    heading_stack: list[str] = []
    table_documents: list[LoadedDocument] = []

    for block in _iter_docx_blocks(doc):
        if isinstance(block, DocxParagraph):
            text = normalize_document_text(block.text)
            if not text:
                continue
            level = _docx_heading_level(block)
            if level is not None:
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(text)
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)
            continue

        markdown = docx_table_to_markdown(block)
        if not markdown:
            continue
        table_documents.append(
            LoadedDocument(
                text=markdown,
                source=str(path),
                metadata={
                    "file_type": ".docx",
                    "content_type": "table",
                    "table_index": len(table_documents),
                    "headings": " > ".join(heading_stack),
                },
            )
        )

    documents: list[LoadedDocument] = []
    text = normalize_document_text("\n\n".join(paragraphs))
    if text:
        documents.append(
            LoadedDocument(
                text=text,
                source=str(path),
                metadata={
                    "file_type": ".docx",
                    "content_type": "text",
                    "headings": " > ".join(heading_stack),
                },
            )
        )

    documents.extend(table_documents)

    return documents


def _load_html(path: Path) -> list[LoadedDocument]:
    soup = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")
    for element in soup(["script", "style", "noscript", "nav", "header", "footer", "aside"]):
        element.decompose()

    content = _semantic_html_container(soup)
    heading_stack = _html_heading_stack(content)
    table_blocks = extract_html_tables(content, {"headings": " > ".join(heading_stack)})
    for table in content.find_all("table"):
        table.decompose()

    normalized = _html_text_to_markdown(content)
    documents: list[LoadedDocument] = []
    if not normalized:
        documents = []
    else:
        documents.append(
            LoadedDocument(
                text=normalized,
                source=str(path),
                metadata={
                    "file_type": path.suffix.lower(),
                    "content_type": "text",
                    "headings": " > ".join(heading_stack),
                },
            )
        )

    for table in table_blocks:
        metadata = {"file_type": path.suffix.lower(), **table.metadata}
        documents.append(LoadedDocument(text=table.markdown, source=str(path), metadata=metadata))

    return documents


def _docx_heading_level(paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if not style_name.startswith("Heading "):
        return None
    try:
        return max(1, min(6, int(style_name.removeprefix("Heading "))))
    except ValueError:
        return None


def _iter_docx_blocks(document: DocxDocumentType):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, document)


def _semantic_html_container(soup: BeautifulSoup):
    return soup.find("main") or soup.find("article") or soup.find("body") or soup


def _html_heading_stack(container) -> list[str]:
    headings = []
    for heading in container.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        text = normalize_document_text(heading.get_text(separator=" "))
        if text:
            headings.append(text)
    return headings


def _html_text_to_markdown(container) -> str:
    lines: list[str] = []
    for element in container.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"], recursive=True):
        text = normalize_document_text(element.get_text(separator=" "))
        if not text:
            continue
        if element.name and element.name.startswith("h"):
            level = int(element.name[1])
            lines.append(f"{'#' * level} {text}")
        elif element.name == "li":
            lines.append(f"- {text}")
        else:
            lines.append(text)
    if not lines:
        lines = [
            line
            for line in normalize_document_text(container.get_text(separator="\n")).splitlines()
            if line.strip()
        ]
    return normalize_document_text("\n\n".join(lines))
