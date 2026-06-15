from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.ingestion.chunker import chunk_documents
from app.ingestion.loaders import LoadedDocument, load_document
from app.ingestion.table_extractors import TableBlock, rows_to_markdown


def test_rows_to_markdown_escapes_cells_and_uses_first_row_as_header():
    markdown = rows_to_markdown(
        [
            ["Metric", "Value"],
            ["Revenue | net", "12\nmillion"],
        ]
    )

    assert markdown == "| Metric | Value |\n| --- | --- |\n| Revenue \\| net | 12 million |"


def test_docx_loader_extracts_tables_as_markdown_documents(tmp_path: Path):
    path = tmp_path / "report.docx"
    doc = DocxDocument()
    doc.add_heading("年度报告", level=1)
    doc.add_paragraph("以下是核心指标。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "营收"
    table.cell(1, 1).text = "100"
    doc.save(path)

    loaded = load_document(path)

    text_doc = next(document for document in loaded if document.metadata["content_type"] == "text")
    table_doc = next(document for document in loaded if document.metadata["content_type"] == "table")
    assert "# 年度报告" in text_doc.text
    assert "以下是核心指标。" in text_doc.text
    assert table_doc.text == "| 指标 | 数值 |\n| --- | --- |\n| 营收 | 100 |"
    assert table_doc.metadata["headings"] == "年度报告"
    assert table_doc.metadata["table_index"] == 0


def test_html_loader_prefers_semantic_content_and_extracts_tables(tmp_path: Path):
    path = tmp_path / "guide.html"
    path.write_text(
        """
        <html>
          <body>
            <nav>导航噪声</nav>
            <main>
              <h1>系统指南</h1>
              <p>正文内容。</p>
              <table>
                <tr><th>模块</th><th>能力</th></tr>
                <tr><td>RAG</td><td>问答</td></tr>
              </table>
            </main>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    loaded = load_document(path)

    combined = "\n".join(document.text for document in loaded)
    table_doc = next(document for document in loaded if document.metadata["content_type"] == "table")
    assert "导航噪声" not in combined
    assert "# 系统指南" in combined
    assert table_doc.text == "| 模块 | 能力 |\n| --- | --- |\n| RAG | 问答 |"
    assert table_doc.metadata["headings"] == "系统指南"


def test_pdf_loader_adds_optional_table_blocks_without_requiring_page_text(tmp_path: Path, monkeypatch):
    path = tmp_path / "table.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=200)
    with path.open("wb") as file:
        writer.write(file)

    monkeypatch.setattr(
        "app.ingestion.loaders.extract_pdf_tables",
        lambda _path: [
            TableBlock(
                markdown="| Metric | Value |\n| --- | --- |\n| Revenue | 100 |",
                metadata={"page": 1, "table_index": 0, "content_type": "table"},
            )
        ],
    )

    loaded = load_document(path)

    assert len(loaded) == 1
    assert loaded[0].text.startswith("| Metric | Value |")
    assert loaded[0].metadata["file_type"] == ".pdf"
    assert loaded[0].metadata["content_type"] == "table"
    assert loaded[0].metadata["page"] == 1


def test_table_chunks_keep_small_tables_whole():
    document = LoadedDocument(
        text="| 指标 | 数值 |\n| --- | --- |\n| 营收 | 100 |\n| 利润 | 20 |",
        source="report.md",
        metadata={"file_type": ".md", "content_type": "table"},
    )

    chunks = chunk_documents([document], chunk_size=120, chunk_overlap=0)

    assert len(chunks) == 1
    assert chunks[0].content.count("| 营收 | 100 |") == 1
    assert chunks[0].metadata["content_type"] == "table"


def test_large_table_chunks_repeat_header_and_do_not_split_rows():
    document = LoadedDocument(
        text=(
            "| 指标 | 数值 |\n"
            "| --- | --- |\n"
            "| 一 | 1111111111 |\n"
            "| 二 | 2222222222 |\n"
            "| 三 | 3333333333 |"
        ),
        source="report.md",
        metadata={"file_type": ".md", "content_type": "table"},
    )

    chunks = chunk_documents([document], chunk_size=55, chunk_overlap=0)

    assert len(chunks) > 1
    assert all(chunk.content.startswith("| 指标 | 数值 |\n| --- | --- |") for chunk in chunks)
    assert any("| 一 | 1111111111 |" in chunk.content for chunk in chunks)
    assert any("| 二 | 2222222222 |" in chunk.content for chunk in chunks)
    assert all(chunk.metadata["content_type"] == "table" for chunk in chunks)
