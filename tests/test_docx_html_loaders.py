from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.loaders import load_document


def test_load_docx_document(tmp_path: Path):
    path = tmp_path / "policy.docx"
    doc = DocxDocument()
    doc.add_heading("制度文档", level=1)
    doc.add_paragraph("员工可以通过知识库查询规章制度。")
    doc.save(path)

    loaded = load_document(path)

    assert len(loaded) == 1
    assert "制度文档" in loaded[0].text
    assert "员工可以通过知识库查询规章制度。" in loaded[0].text
    assert loaded[0].metadata["file_type"] == ".docx"


def test_load_html_document_removes_scripts_and_keeps_text(tmp_path: Path):
    path = tmp_path / "guide.html"
    path.write_text(
        """
        <html>
          <head><script>console.log("ignore me")</script></head>
          <body>
            <h1>系统指南</h1>
            <p>RAG 系统支持知识库问答。</p>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    loaded = load_document(path)

    assert len(loaded) == 1
    assert "系统指南" in loaded[0].text
    assert "RAG 系统支持知识库问答。" in loaded[0].text
    assert "console.log" not in loaded[0].text
    assert loaded[0].metadata["file_type"] == ".html"
